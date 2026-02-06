"""
Document Scraper - PDF, Excel, CSV extraction

Extracts structured data from document files.
"""
import logging
import asyncio
from typing import Dict, Any, List, Optional
from pathlib import Path
import tempfile
import httpx

from app.scraper.logic.base import BaseScraper
from app.schemas import ScrapeResult, ScrapeFailureReason, DocumentConfig

logger = logging.getLogger(__name__)


class DocumentScraper(BaseScraper):
    """
    Extract data from PDF, Excel, CSV, and Word documents.
    """
    
    def get_name(self) -> str:
        return "document"
    
    def can_handle(self, url: str) -> bool:
        """Check if URL points to a document"""
        doc_extensions = ['.pdf', '.xlsx', '.xls', '.csv', '.docx', '.doc']
        return any(url.lower().endswith(ext) for ext in doc_extensions)
    
    async def scrape(
        self,
        url: str,
        schema: Dict[str, Any],
        job_id: str,
        **kwargs
    ) -> ScrapeResult:
        """
        Extract data from document
        
        Args:
            url: Document URL
            schema: Extraction schema (not used for tables)
            job_id: Job identifier
            **kwargs: document_config (DocumentConfig)
        """
        logger.info(f"Starting document extraction for {url}")
        
        doc_config: DocumentConfig = kwargs.get('document_config', DocumentConfig())
        
        try:
            # Download document
            file_path = await self._download_document(url)
            
            # Determine file type
            file_ext = Path(url).suffix.lower()
            
            # Extract based on type
            if file_ext == '.pdf':
                data = await self._extract_pdf(file_path, doc_config)
            elif file_ext in ['.xlsx', '.xls']:
                data = await self._extract_excel(file_path, doc_config)
            elif file_ext == '.csv':
                data = await self._extract_csv(file_path)
            elif file_ext in ['.docx', '.doc']:
                data = await self._extract_word(file_path, doc_config)
            else:
                return self.failure(
                    reason=ScrapeFailureReason.VALIDATION_FAILED,
                    message=f"Unsupported document type: {file_ext}"
                )
            
            # Cleanup temp file
            Path(file_path).unlink(missing_ok=True)
            
            return ScrapeResult(
                success=True,
                status="success",
                data=data,
                strategy_used=self.get_name(),
                confidence=0.90,
                metadata={
                    "document_type": file_ext,
                    "source_url": url
                }
            )
            
        except Exception as e:
            logger.error(f"Document extraction failed: {e}", exc_info=True)
            return self.failure(
                reason=ScrapeFailureReason.UNKNOWN,
                message=f"Document extraction error: {str(e)}",
                errors=[str(e)]
            )
    
    async def _download_document(self, url: str) -> str:
        """Download document to temp file"""
        async with httpx.AsyncClient(timeout=60) as client:
            response = await client.get(url, follow_redirects=True)
            response.raise_for_status()
            
            # Save to temp file
            suffix = Path(url).suffix
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(response.content)
                return tmp.name
    
    async def _extract_pdf(
        self,
        file_path: str,
        config: DocumentConfig
    ) -> Dict[str, Any]:
        """Extract data from PDF"""
        try:
            import pdfplumber
        except ImportError:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
        
        extracted = {
            "text": "",
            "tables": [],
            "metadata": {}
        }
        
        with pdfplumber.open(file_path) as pdf:
            extracted["metadata"] = {
                "pages": len(pdf.pages),
                "info": pdf.metadata
            }
            
            # Determine page range
            pages_to_process = self._get_page_range(
                len(pdf.pages),
                config.page_range
            )
            
            for page_num in pages_to_process:
                page = pdf.pages[page_num]
                
                # Extract text
                if config.extract_text:
                    text = page.extract_text()
                    if text:
                        extracted["text"] += f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Extract tables
                if config.extract_tables:
                    tables = page.extract_tables()
                    for table in tables:
                        if table:
                            extracted["tables"].append({
                                "page": page_num + 1,
                                "data": table
                            })
        
        return extracted
    
    async def _extract_excel(
        self,
        file_path: str,
        config: DocumentConfig
    ) -> Dict[str, Any]:
        """Extract data from Excel"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas not installed. Run: pip install pandas openpyxl")
        
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        
        extracted = {
            "sheets": {},
            "metadata": {
                "sheet_names": excel_file.sheet_names
            }
        }
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            
            # Convert to records
            extracted["sheets"][sheet_name] = {
                "rows": len(df),
                "columns": list(df.columns),
                "data": df.to_dict('records')
            }
        
        return extracted
    
    async def _extract_csv(self, file_path: str) -> Dict[str, Any]:
        """Extract data from CSV"""
        try:
            import pandas as pd
        except ImportError:
            raise ImportError("pandas not installed. Run: pip install pandas")
        
        df = pd.read_csv(file_path)
        
        return {
            "rows": len(df),
            "columns": list(df.columns),
            "data": df.to_dict('records')
        }
    
    async def _extract_word(
        self,
        file_path: str,
        config: DocumentConfig
    ) -> Dict[str, Any]:
        """Extract data from Word document"""
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        doc = Document(file_path)
        
        extracted = {
            "text": "",
            "tables": [],
            "metadata": {
                "paragraphs": len(doc.paragraphs)
            }
        }
        
        # Extract text
        if config.extract_text:
            extracted["text"] = "\n".join([p.text for p in doc.paragraphs])
        
        # Extract tables
        if config.extract_tables:
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text for cell in row.cells])
                extracted["tables"].append(table_data)
        
        return extracted
    
    def _get_page_range(
        self,
        total_pages: int,
        page_range: Optional[str]
    ) -> List[int]:
        """
        Parse page range string
        
        Examples:
            "1-5" -> [0, 1, 2, 3, 4]
            "all" -> [0, 1, ..., total_pages-1]
            "1,3,5" -> [0, 2, 4]
        """
        if not page_range or page_range.lower() == "all":
            return list(range(total_pages))
        
        pages = []
        
        # Handle comma-separated pages
        if ',' in page_range:
            for part in page_range.split(','):
                page_num = int(part.strip()) - 1  # Convert to 0-indexed
                if 0 <= page_num < total_pages:
                    pages.append(page_num)
        
        # Handle range
        elif '-' in page_range:
            start, end = page_range.split('-')
            start_page = int(start.strip()) - 1
            end_page = int(end.strip())
            pages = list(range(
                max(0, start_page),
                min(total_pages, end_page)
            ))
        
        # Single page
        else:
            page_num = int(page_range) - 1
            if 0 <= page_num < total_pages:
                pages = [page_num]
        
        return pages
    
    async def extract(
        self,
        document_url: str,
        extract_tables: bool = True,
        extract_text: bool = True
    ) -> ScrapeResult:
        """
        Public method for document extraction
        
        Convenience wrapper around scrape()
        """
        doc_config = DocumentConfig(
            extract_tables=extract_tables,
            extract_text=extract_text
        )
        
        return await self.scrape(
            url=document_url,
            schema={},
            job_id="document",
            document_config=doc_config
        )
